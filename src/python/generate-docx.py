#!/usr/bin/env python3
# usage: python src/python/generate-docx.py requirements_doc .
import sys
import os
import urllib.request, json 
import docx
from docx import Document
from docx.shared import Inches


table, path, app, db, user, ht_protocol, web_host, web_port = '', '', '', '', '', '', '', ''
def main():
    set_vars()
    tableData = getTableData()
    buildDoc(table,tableData)
    sys.exit()


def set_vars():
    try:
        global table, path, app, db, user, ht_protocol, web_host, web_port
        table = sys.argv[1]
        path = str(sys.argv[2] or '.')
        app = os.environ['postgres_db_name']
        db = os.environ['postgres_db_name']
        user= os.environ['postgres_db_user']
        ht_protocol = os.environ['ht_protocol']
        web_host= os.environ['web_host']
        web_port = os.environ['port']
    except(IndexError) as error:
        print (error) 
        traceback.print_stack()
        sys.exit(1)

    return


def getTableData():
    try:
        url = ht_protocol + "://" + web_host + ":" + web_port + "/" + app + "/hiselect/" + table + "?pg-size=10000&bid=0"
        with urllib.request.urlopen(url) as url:
            data = json.loads(url.read().decode())
            print(data)
    except (Exception) as error:
        print(error)
    finally:
        return data['dat']


def buildDoc(table,tableData):
    doc = Document()
    #debug print('Loaded {}'.format(tableData))
    for row in tableData:
        print ('row {}'.format(row))
        level = 'Heading ' + str(row['level']+1)
        title = str(row['logical_order'] or '') + " " + row['name']
        if row['id'] == 0:
            doc.add_heading(title, 0)
            doc.add_section()
            doc.add_section()
        else:
            heading = doc.add_heading( title , row['level'])
            heading.style = doc.styles[level]
            p = doc.add_paragraph(row['description'])
            #if row['img_relative_path']:
            #    ip = document.add_paragraph()
            #    r = p.add_run()
            #    r.add_text(row['img_name'])
            #    r.add_picture('/tmp/foo.jpg')
            if row['src']:
                src = doc.add_paragraph(row['src'])
                src.style = doc.styles['Body Text 3']

    full_path=path + "/" + table.replace('_','.').replace('doc','docx')    
    doc.save(full_path)
    return 0


main()
