#!/usr/bin/env python3
# usage: python src/python/append-to-docx.py requirements_doc /path/to/file.docx
import sys
import os
import urllib.request, json 
import docx
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import StringIO



ikey, full_path, table, path, app, db, user, ht_protocol, web_host, web_port = 0, '', '', '', '', '', '', '', '', ''
def main():
    set_vars()
    tableData, met = getTableData(None)
    buildDoc(table,tableData)
    sys.exit()


def set_vars():
    try:
        global full_path, table, path, app, db, user, ht_protocol, web_host, web_port
        table = sys.argv[1]
        full_path = str(sys.argv[2] or '.')
        app = os.environ['postgres_db_name']
        db = os.environ['postgres_db_name']
        user= os.environ['postgres_db_user']
        ht_protocol = os.environ['ht_protocol']
        web_host= os.environ['web_host']
        web_port = os.environ['port']
    except(IndexError) as error:
        print (error) 
        traceback.print_stack()
        sys.exit(1)

    return


def getTableData(url):
       
    if ( url == None):
        url = ht_protocol + "://" + web_host + ":" + web_port + "/" + app + "/hiselect/" + table + "?pg-size=10000&bid=0"
    try:
        with urllib.request.urlopen(url) as url:
            data = json.loads(url.read().decode())
            # print(data['dat'])
            print(data['met']['meta_cols'])
            return data['dat'], data['met']['meta_cols']
    except (Exception) as error:
        print(error)


def buildDoc(table,tableData):
    doc = Document(full_path) # open an existing document with existing styles
    #debug print('Loaded {}'.format(tableData))
    for row in tableData:
        # print ('row {}'.format(row))
        level = row['level']
        if ( level == 0 ):
            continue
        if ( level == 1 ):
            levelStyle = 'style-kone-blue-heading-01'
        else: 
            levelStyle = 'Heading ' + str(level)
        title = row['name']
        heading = doc.add_heading( title , level)
        heading.style = doc.styles[levelStyle]
        p = doc.add_paragraph(row['description'])
        if row['src']:
            src = doc.add_paragraph(row['src'])
            src.style = doc.styles['Body Text 3']
        if ( row['formats'] ):
            if (not( str(row['formats']).strip() == "")):
                print ("formats: " + row['formats'])
                formats = json.loads(row['formats'])
                listing_url = formats['listing-url']
                print ("listing_url: " + listing_url)
                listingData , met = getTableData(listing_url)
                #for style in doc.styles:
                    #print ( style )
                global ikey
                doc_table = doc.add_table(1, (len(listingData[0].keys())-2), doc.styles['Table Grid'])
                for lrow in listingData:
                    mikey = -1
                    ikey = 0
                    for mlrowkey in sorted (met.keys()):
                        mlrow = met[str(mlrowkey)]
                        key = mlrow['attribute_name']
                        mikey = mikey + 1
                        if ( not ( key in lrow.keys() ) ):
                            continue
                        if ( not ( key == 'id' or key == 'guid' )): 
                            if ( ikey == 0 ):
                                cells = doc_table.add_row().cells
                                thkey = 0
                                for th in sorted (met.keys()): # this is the title row
                                    mthrow = met[str(th)]
                                    ky = mthrow['attribute_name']
                                    print ( "lrow.keys:")
                                    print ( "ky: " + str(ky))
                                    print(lrow.keys())
                                    if ( (not ( ky in lrow.keys() )) or ky == 'id' or ky == 'guid'):
                                        continue
                                    else:
                                        cells[thkey].text = str(ky)
                                        thkey=thkey+1
                            #print ( str(lrow[key]) ) # the value cells[0].text = lrow[key]
                            cells[ikey].text = str(lrow[key])
                            ikey = ikey + 1
        if row['img_relative_path']:
            ip = doc.add_paragraph()
            r = ip.add_run()
            r.add_text(row['img_name'])
            r.add_text("\n")
            r.add_picture(row['img_relative_path'], width = Cm(15.0))

    doc.save(full_path)
    return 0


main()
